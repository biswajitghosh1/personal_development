from flask import (Flask, render_template, request, redirect, url_for,
                   flash, jsonify, Response, send_file)
from datetime import date
from dateutil.parser import parse as parse_date
from io import BytesIO, StringIO
import csv, re, os
from collections import Counter
from models import db, Entry, User
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash
from docx import Document
from docx.shared import RGBColor

def create_app(test_config=None):
    app = Flask(__name__, instance_relative_config=True, template_folder="templates", static_folder="static")
    app.config.from_mapping(
        SECRET_KEY=os.environ.get("DJ_SECRET", "dev-secret"),
        SQLALCHEMY_DATABASE_URI=f"sqlite:///{os.path.join(app.instance_path, 'journal.db')}",
        SQLALCHEMY_TRACK_MODIFICATIONS=False,
    )
    if test_config:
        app.config.update(test_config)
    try:
        os.makedirs(app.instance_path, exist_ok=True)
    except OSError:
        pass

    db.init_app(app)

    login_manager = LoginManager()
    login_manager.login_view = "login"
    login_manager.init_app(app)

    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))

    @app.context_processor
    def inject_now():
        from datetime import datetime
        return {"now": datetime.utcnow}

    @app.route("/")
    def index():
        q_date = request.args.get("date")
        if q_date:
            try:
                dt = parse_date(q_date).date()
            except Exception:
                flash("Invalid date format", "error")
                return redirect(url_for("index"))
            entries = Entry.query.filter_by(entry_date=dt).order_by(Entry.entry_date.desc(), Entry.created_at.desc()).all()
        else:
            entries = Entry.query.order_by(Entry.entry_date.desc(), Entry.created_at.desc()).limit(200).all()
        return render_template("index.html", entries=entries)

    @app.route("/login", methods=["GET","POST"])
    def login():
        if request.method == "POST":
            username = request.form.get("username","").strip()
            password = request.form.get("password","").strip()
            user = User.query.filter_by(username=username).first()
            if user and user.check_password(password):
                login_user(user)
                flash("Logged in", "success")
                return redirect(url_for("index"))
            flash("Invalid username or password", "error")
            return redirect(url_for("login"))
        return render_template("login.html")

    @app.route("/logout")
    @login_required
    def logout():
        logout_user()
        flash("Logged out", "success")
        return redirect(url_for("index"))

    @app.route("/add", methods=["GET","POST"])
    @login_required
    def add_entry():
        if request.method == "POST":
            entry_date = request.form.get("entry_date")
            title = request.form.get("title","").strip()
            content = request.form.get("content","").strip()
            important = bool(request.form.get("important"))
            if not entry_date or not title or not content:
                flash("Date, title and content are required", "error")
                return redirect(url_for("add_entry"))
            try:
                dt = parse_date(entry_date).date()
            except Exception:
                flash("Invalid date", "error")
                return redirect(url_for("add_entry"))
            e = Entry(entry_date=dt, title=title, content=content, important=important)
            db.session.add(e)
            db.session.commit()
            flash("Entry added", "success")
            return redirect(url_for("index"))
        default_date = date.today().isoformat()
        return render_template("add_edit.html", entry=None, default_date=default_date)

    @app.route("/edit/<int:entry_id>", methods=["GET","POST"])
    @login_required
    def edit_entry(entry_id):
        e = Entry.query.get_or_404(entry_id)
        if request.method == "POST":
            entry_date = request.form.get("entry_date")
            title = request.form.get("title","").strip()
            content = request.form.get("content","").strip()
            important = bool(request.form.get("important"))
            if not entry_date or not title or not content:
                flash("Date, title and content are required", "error")
                return redirect(url_for("edit_entry", entry_id=entry_id))
            try:
                dt = parse_date(entry_date).date()
            except Exception:
                flash("Invalid date", "error")
                return redirect(url_for("edit_entry", entry_id=entry_id))
            e.entry_date = dt
            e.title = title
            e.content = content
            e.important = important
            db.session.commit()
            flash("Entry updated", "success")
            return redirect(url_for("index"))
        return render_template("add_edit.html", entry=e, default_date=e.entry_date.isoformat())

    @app.route("/delete/<int:entry_id>", methods=["POST"])
    @login_required
    def delete_entry(entry_id):
        e = Entry.query.get_or_404(entry_id)
        db.session.delete(e)
        db.session.commit()
        flash("Entry deleted", "success")
        return redirect(url_for("index"))

    def basic_summary_for_entries(entries):
        total_entries = len(entries)
        total_words = 0
        counter = Counter()
        word_re = re.compile(r"\b[a-zA-Z']+\b")
        stopwords = {
            "the","and","a","to","of","in","is","it","for","on","that","with","as","this","are","was","but","be","or","by","an","from"
        }
        for e in entries:
            words = word_re.findall(e.title.lower() + " " + e.content.lower())
            filtered = [w for w in words if w not in stopwords and len(w) > 1]
            total_words += len(filtered)
            counter.update(filtered)
        top = counter.most_common(10)
        return {
            "total_entries": total_entries,
            "total_words": total_words,
            "average_words_per_entry": total_words / total_entries if total_entries else 0,
            "top_words": top,
        }

    @app.route("/summary/<int:year>/<int:month>", methods=["GET"])
    def monthly_summary(year, month):
        start = date(year, month, 1)
        if month == 12:
            end = date(year + 1, 1, 1)
        else:
            end = date(year, month + 1, 1)
        entries = Entry.query.filter(Entry.entry_date >= start, Entry.entry_date < end).order_by(Entry.entry_date.asc()).all()
        summary = basic_summary_for_entries(entries)
        return render_template("summary.html", year=year, month=month, entries=entries, summary=summary)

    @app.route("/export", methods=["GET"])
    @login_required
    def export():
        fmt = request.args.get("format", "csv").lower()
        start = request.args.get("start")
        end = request.args.get("end")
        try:
            s = parse_date(start).date() if start else None
            e = parse_date(end).date() if end else None
        except Exception:
            return "Invalid date", 400

        q = Entry.query
        if s:
            q = q.filter(Entry.entry_date >= s)
        if e:
            q = q.filter(Entry.entry_date <= e)
        results = q.order_by(Entry.entry_date.asc()).all()

        if fmt == "json":
            return jsonify([r.to_dict() for r in results])

        if fmt == "docx":
            doc = Document()
            doc.add_heading('Journal Export', level=1)
            doc.add_paragraph(f"Range: {start or 'start'} — {end or 'end'}")
            for r in results:
                p = doc.add_paragraph()
                run_text = f"{r.entry_date.isoformat()} — {r.title}"
                if r.important:
                    run_text = "[IMPORTANT] " + run_text
                run = p.add_run(run_text)
                run.bold = True
                if r.important:
                    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
                doc.add_paragraph(r.content)
                doc.add_paragraph("")
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            return send_file(
                bio,
                as_attachment=True,
                download_name=f"journal_{start or 'all'}_{end or 'all'}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # CSV fallback
        si = StringIO()
        cw = csv.writer(si)
        cw.writerow(["id","entry_date","title","important","content","created_at","updated_at"])
        for r in results:
            cw.writerow([r.id, r.entry_date.isoformat(), r.title, r.important, r.content, r.created_at.isoformat() if r.created_at else "", r.updated_at.isoformat() if r.updated_at else ""])
        output = si.getvalue()
        return Response(
            output,
            mimetype="text/csv",
            headers={"Content-Disposition": f"attachment; filename=entries_{start or 'all'}_{end or 'all'}.csv"}
        )

    @app.route("/api/entries", methods=["GET"])
    @login_required
    def api_entries():
        start = request.args.get("start")
        end = request.args.get("end")
        limit = int(request.args.get("limit", "100"))
        offset = int(request.args.get("offset", "0"))
        try:
            s = parse_date(start).date() if start else None
            e = parse_date(end).date() if end else None
        except Exception:
            return jsonify({"error":"invalid date"}), 400
        q = Entry.query
        if s:
            q = q.filter(Entry.entry_date >= s)
        if e:
            q = q.filter(Entry.entry_date <= e)
        items = q.order_by(Entry.entry_date.asc()).limit(limit).offset(offset).all()
        return jsonify([i.to_dict() for i in items])

    return app

if __name__ == "__main__":
    application = create_app()
    application.run(debug=True)
